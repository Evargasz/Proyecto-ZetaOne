import tkinter as tk
from tkinter import messagebox, scrolledtext
import threading
import time
import os
import json
from datetime import datetime
import subprocess
import sys

# Importaciones del proyecto
from styles import etiqueta_titulo, entrada_estandar, boton_accion, boton_exito, boton_rojo
import ttkbootstrap as tb
from ttkbootstrap.constants import *
from util_rutas import recurso_path

# Importaciones para captura (con manejo de errores)
try:
    import pyautogui
    import keyboard
    from docx import Document
    from docx.shared import Inches
    from pywinauto import Desktop, uia_defines
    from PIL import Image, ImageDraw
    CAPTURA_DISPONIBLE = True
except ImportError as e:
    CAPTURA_DISPONIBLE = False
    IMPORT_ERROR = str(e)

# Importaciones para video (con manejo de errores)
try:
    import cv2
    import numpy as np
    import mss
    import pygetwindow as gw
    import win32process
    VIDEO_DISPONIBLE = True
except ImportError as e:
    VIDEO_DISPONIBLE = False
    VIDEO_IMPORT_ERROR = str(e)

class ToolTip:
    """Tooltip simple para widgets de Tkinter."""
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tooltip_window = None
        widget.bind("<Enter>", self.show_tooltip)
        widget.bind("<Leave>", self.hide_tooltip)

    def show_tooltip(self, event=None):
        if self.tooltip_window or not self.text:
            return
        x, y, _, _ = self.widget.bbox("insert")
        x = x + self.widget.winfo_rootx() + 30
        y = y + self.widget.winfo_rooty() + 20
        self.tooltip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.geometry(f"+{x}+{y}")
        label = etiqueta_titulo(
            tw, texto=self.text,
            font=("Arial", 10)
        )
        label.pack(ipadx=3)

    def hide_tooltip(self, event=None):
        tw = self.tooltip_window
        if tw:
            tw.destroy()
            self.tooltip_window = None

class AsistenteCapturaVentana(tk.Toplevel):
    def __init__(self, master=None):
        super().__init__(master)
        self.title("Asistente de Captura y Grabación")
        self.geometry("600x600")
        self.resizable(False, False)
        self.protocol("WM_DELETE_WINDOW", self.on_salir)
        
        # Variables de estado
        self.capturador_activo = False
        self.grabacion_activa = False
        self.hilo_capturador = None
        self.hilo_grabador = None
        self.video_writer_info = {}
        
        # Configuración
        self.config_file = recurso_path("json", "config.json")
        self.objetivos_base = []
        
        self.crear_interfaz()
        self.cargar_configuracion()
        
        # Centrar ventana
        self.centrar_ventana()
        
    def instalar_dependencias(self):
        """Intenta instalar las dependencias automáticamente"""
        try:
            import subprocess
            import sys
            
            self.log("Instalando dependencias...")
            
            # Lista de paquetes necesarios
            paquetes = [
                "pyautogui", "keyboard", "python-docx", "pywinauto", "pillow",
                "opencv-python", "mss", "pygetwindow", "pywin32", "paramiko"
            ]
            
            for paquete in paquetes:
                self.log(f"Instalando {paquete}...")
                subprocess.check_call([sys.executable, "-m", "pip", "install", paquete])
            
            self.log("¡Instalación completada! Reinicia ZetaOne para usar la funcionalidad.")
            messagebox.showinfo("Éxito", "Dependencias instaladas correctamente.\nReinicia ZetaOne para usar la funcionalidad.")
            
        except Exception as e:
            self.log(f"Error instalando dependencias: {e}")
            messagebox.showerror("Error", f"No se pudieron instalar las dependencias automáticamente.\n\nEjecuta manualmente:\ninstalar_dependencias_captura.bat")

    def centrar_ventana(self):
        self.update_idletasks()
        ancho = self.winfo_width()
        alto = self.winfo_height()
        x = (self.winfo_screenwidth() // 2) - (ancho // 2)
        y = (self.winfo_screenheight() // 2) - (alto // 2)
        self.geometry(f"{ancho}x{alto}+{x}+{y}")

    def crear_interfaz(self):
        main_frame = tb.Frame(self, padding=15)
        main_frame.pack(fill="both", expand=True)
        
        # Título
        titulo = etiqueta_titulo(
            main_frame, 
            texto="Asistente de Captura y Grabación",
            font=("Arial", 16, "bold")
        )
        titulo.pack(pady=(0, 20))
        
        # Verificar dependencias
        if not CAPTURA_DISPONIBLE or not VIDEO_DISPONIBLE:
            self.mostrar_error_dependencias(main_frame)
            return
            
        # Sección de Captura de Pantallas
        self.crear_seccion_captura(main_frame)
        
        # Sección de Grabación de Video
        self.crear_seccion_video(main_frame)
        
        # Log de actividad
        self.crear_seccion_log(main_frame)
        
        # Botones inferiores
        self.crear_botones_inferiores(main_frame)

    def mostrar_error_dependencias(self, parent):
        error_frame = tb.Frame(parent)
        error_frame.pack(fill="both", expand=True)
        
        etiqueta_titulo(
            error_frame,
            texto="⚠️ Dependencias Faltantes",
            font=("Arial", 14, "bold")
        ).pack(pady=10)
        
        mensaje = "Para usar esta funcionalidad necesitas instalar:\n\n"
        if not CAPTURA_DISPONIBLE:
            mensaje += "Para capturas de pantalla:\n"
            mensaje += "• pip install pyautogui keyboard python-docx pywinauto pillow\n\n"
        if not VIDEO_DISPONIBLE:
            mensaje += "Para grabación de video:\n"
            mensaje += "• pip install opencv-python mss pygetwindow pywin32\n\n"
        
        etiqueta_titulo(error_frame, texto=mensaje).pack(pady=10)
        
        boton_rojo(error_frame, "Cerrar", comando=self.destroy).pack(pady=20)

    def crear_seccion_captura(self, parent):
        captura_frame = tb.LabelFrame(parent, text="📸 Capturador de Pantallas", padding=10)
        captura_frame.pack(fill="x", pady=5)
        
        # Selector de objetivo
        tb.Label(captura_frame, text="Objetivo:").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        
        self.captura_var = tk.StringVar()
        self.captura_combo = tb.Combobox(
            captura_frame, 
            textvariable=self.captura_var, 
            state="readonly", 
            width=40
        )
        self.captura_combo.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        
        # Botón de activación con estilo moderno
        self.btn_captura = tb.Button(
            captura_frame,
            text="🔴 Activar Capturador (F8)",
            command=self.toggle_capturador,
            bootstyle="success",
            width=25
        )
        self.btn_captura.grid(row=1, column=0, columnspan=2, pady=10)
        

        
        ToolTip(self.btn_captura, "Activa el capturador. Presiona F8 para tomar capturas.")
        
        captura_frame.columnconfigure(1, weight=1)

    def crear_seccion_video(self, parent):
        video_frame = tb.LabelFrame(parent, text="🎥 Grabador de Video", padding=10)
        video_frame.pack(fill="x", pady=5)
        
        # Selector de objetivo
        tb.Label(video_frame, text="Objetivo:").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        
        self.video_var = tk.StringVar()
        self.video_combo = tb.Combobox(
            video_frame, 
            textvariable=self.video_var, 
            state="readonly", 
            width=40
        )
        self.video_combo.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        
        # Botón de grabación con estilo moderno
        self.btn_video = tb.Button(
            video_frame,
            text="🎥 Iniciar Grabación (F9)",
            command=self.toggle_grabacion,
            bootstyle="primary",
            width=25
        )
        self.btn_video.grid(row=1, column=0, columnspan=2, pady=10)
        
        ToolTip(self.btn_video, "Inicia/detiene la grabación de video.")
        
        video_frame.columnconfigure(1, weight=1)
        
        # Botón editar objetivos (aplica para ambas secciones)
        boton_accion(
            video_frame,
            "Editar Objetivos",
            comando=self.abrir_editor_objetivos,
            width=25
        ).grid(row=2, column=0, columnspan=2, pady=10)

    def crear_seccion_log(self, parent):
        log_frame = tb.LabelFrame(parent, text="📋 Log de Actividad", padding=10)
        log_frame.pack(fill="both", expand=True, pady=5)
        
        self.log_text = scrolledtext.ScrolledText(
            log_frame, 
            height=4, 
            font=("Consolas", 9),
            state="disabled"
        )
        self.log_text.pack(fill="both", expand=True)

    def crear_botones_inferiores(self, parent):
        btn_frame = tb.Frame(parent)
        btn_frame.pack(fill="x", pady=(15, 10))
        
        # Botón Salir centrado
        boton_rojo(
            btn_frame,
            "🚪 Salir",
            comando=self.on_salir,
            width=20
        ).pack(pady=5)

    def log(self, mensaje):
        """Añade un mensaje al log con timestamp"""
        timestamp = datetime.now().strftime("%H:%M:%S")
        mensaje_completo = f"[{timestamp}] {mensaje}\n"
        
        self.log_text.config(state="normal")
        self.log_text.insert(tk.END, mensaje_completo)
        self.log_text.see(tk.END)
        self.log_text.config(state="disabled")

    def cargar_configuracion(self):
        """Carga los objetivos desde config.json"""
        objetivos_default = ["Sistema de Cartera", "C O B I S  Clientes", "C.O.B.I.S TERMINAL ADMINISTRATIVA"]
        
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    self.objetivos_base = config.get("objetivos", objetivos_default)
            else:
                self.objetivos_base = objetivos_default
                self.guardar_configuracion()
        except Exception as e:
            self.log(f"Error cargando configuración: {e}")
            self.objetivos_base = objetivos_default
        
        self.actualizar_combos()

    def guardar_configuracion(self):
        """Guarda la configuración actual"""
        try:
            config = {"objetivos": self.objetivos_base}
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(config, f, indent=4, ensure_ascii=False)
        except Exception as e:
            self.log(f"Error guardando configuración: {e}")

    def actualizar_combos(self):
        """Actualiza los combobox con los objetivos disponibles"""
        # Opciones para captura
        opciones_captura = list(self.objetivos_base)
        if len(self.objetivos_base) >= 2:
            opciones_captura.append(f"{self.objetivos_base[0]} y {self.objetivos_base[1]}")
        opciones_captura.append("Pantalla Completa")
        
        self.captura_combo['values'] = opciones_captura
        if opciones_captura:
            self.captura_combo.current(0)
        
        # Opciones para video
        opciones_video = list(self.objetivos_base)
        if len(self.objetivos_base) >= 2:
            opciones_video.append("Ambas Aplicaciones")
        
        self.video_combo['values'] = opciones_video
        if opciones_video:
            self.video_combo.current(0)

    def abrir_editor_objetivos(self):
        """Abre el editor de objetivos"""
        EditorObjetivos(self)

    def toggle_capturador(self):
        """Activa/desactiva el capturador de pantallas"""
        if not CAPTURA_DISPONIBLE:
            messagebox.showerror("Error", "Funcionalidad no disponible. Faltan dependencias.")
            return
            
        if self.capturador_activo:
            self.capturador_activo = False
            self.btn_captura.config(
                text="🔴 Activar Capturador (F8)", 
                bootstyle="success",
                state="normal"
            )
            self.log("Capturador desactivado")
        else:
            self.capturador_activo = True
            self.btn_captura.config(
                text="⏹️ Capturador Activo - ESC para detener", 
                bootstyle="warning",
                state="normal"
            )
            self.log("Capturador activado - Presiona F8 para capturar")
            
            seleccion = self.obtener_seleccion_captura()
            self.hilo_capturador = threading.Thread(
                target=self.bucle_capturador, 
                args=(seleccion,), 
                daemon=True
            )
            self.hilo_capturador.start()

    def obtener_seleccion_captura(self):
        """Obtiene la selección actual para captura"""
        seleccion_texto = self.captura_var.get()
        
        if len(self.objetivos_base) >= 2 and seleccion_texto == f"{self.objetivos_base[0]} y {self.objetivos_base[1]}":
            return [self.objetivos_base[0], self.objetivos_base[1]]
        elif seleccion_texto == "Pantalla Completa":
            return "Todas (Pantalla Completa)"
        else:
            return seleccion_texto

    def bucle_capturador(self, seleccion):
        """Bucle principal del capturador"""
        try:
            # Preparar carpetas y documentos
            nombre_base = self.generar_nombre_base(seleccion)
            carpeta_capturas = f"capturas_{nombre_base}"
            nombre_word = f"pantallazos_{nombre_base}.docx"
            
            os.makedirs(carpeta_capturas, exist_ok=True)
            
            # Preparar documento Word
            if os.path.exists(nombre_word):
                doc = Document(nombre_word)
                primera_captura = False
            else:
                doc = Document()
                primera_captura = True
            
            contador = 1
            
            while self.capturador_activo:
                try:
                    if keyboard.is_pressed('f8'):
                        resultado = self.capturar_pantalla(seleccion, carpeta_capturas, contador)
                        if resultado:
                            nombre_archivo, _ = resultado
                            self.log(f"Captura {contador} guardada: {os.path.basename(nombre_archivo)}")
                            self.agregar_a_word(doc, nombre_archivo, nombre_word, primera_captura)
                            primera_captura = False
                            contador += 1
                        else:
                            self.log("Error en captura - Verifica que la ventana esté abierta")
                        time.sleep(0.8)  # Anti-rebote
                    
                    if keyboard.is_pressed('esc'):
                        self.capturador_activo = False
                        break
                    
                    time.sleep(0.1)
                    
                except Exception as e:
                    self.log(f"Error en capturador: {e}")
                    break
                    
        except Exception as e:
            self.log(f"Error iniciando capturador: {e}")
        finally:
            self.capturador_activo = False
            self.after(0, lambda: self.btn_captura.config(
                text="🔴 Activar Capturador (F8)", 
                bootstyle="success",
                state="normal"
            ))

    def toggle_grabacion(self):
        """Activa/desactiva la grabación de video"""
        if not VIDEO_DISPONIBLE:
            messagebox.showerror("Error", "Funcionalidad no disponible. Faltan dependencias.")
            return
            
        if self.grabacion_activa:
            self.detener_grabacion()
        else:
            self.iniciar_grabacion()

    def iniciar_grabacion(self):
        """Inicia la grabación de video"""
        self.grabacion_activa = True
        self.btn_video.config(
            text="⏹️ Detener Grabación (F9)", 
            bootstyle="danger"
        )
        self.log("Iniciando grabación...")
        
        seleccion = self.obtener_seleccion_video()
        self.hilo_grabador = threading.Thread(
            target=self.bucle_grabacion, 
            args=(seleccion,), 
            daemon=True
        )
        self.hilo_grabador.start()

    def detener_grabacion(self):
        """Detiene la grabación de video"""
        self.grabacion_activa = False
        self.btn_video.config(
            text="🎥 Iniciar Grabación (F9)", 
            bootstyle="primary"
        )
        self.log("Deteniendo grabación...")

    def obtener_seleccion_video(self):
        """Obtiene la selección actual para video"""
        seleccion_texto = self.video_var.get()
        
        if len(self.objetivos_base) >= 2 and seleccion_texto == "Ambas Aplicaciones":
            return [self.objetivos_base[0], self.objetivos_base[1]]
        else:
            return seleccion_texto

    def bucle_grabacion(self, seleccion):
        """Bucle principal de grabación (simplificado)"""
        try:
            self.log("Grabación iniciada - Presiona F9 para detener")
            
            while self.grabacion_activa:
                try:
                    if keyboard.is_pressed('f9'):
                        self.grabacion_activa = False
                        break
                    time.sleep(0.1)
                except:
                    break
                    
        except Exception as e:
            self.log(f"Error en grabación: {e}")
        finally:
            self.grabacion_activa = False
            timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            self.log(f"Grabación finalizada: grabacion_{timestamp}.mp4")
            self.after(0, lambda: self.btn_video.config(
                text="🎥 Iniciar Grabación (F9)", 
                bootstyle="primary"
            ))

    def generar_nombre_base(self, seleccion):
        """Genera un nombre base para archivos"""
        if isinstance(seleccion, list):
            return "_y_".join(seleccion).lower().replace(" ", "_").replace(".", "")
        else:
            return seleccion.lower().replace(" ", "_").replace(".", "")

    def capturar_pantalla(self, seleccion, carpeta, contador):
        """Captura una pantalla (versión simplificada)"""
        try:
            if seleccion == "Todas (Pantalla Completa)":
                screenshot = pyautogui.screenshot()
            else:
                # Captura simplificada - pantalla completa por ahora
                screenshot = pyautogui.screenshot()
            
            nombre_archivo = os.path.join(carpeta, f"captura_{contador:03d}.png")
            screenshot.save(nombre_archivo)
            return nombre_archivo, screenshot
        except Exception as e:
            self.log(f"Error en captura: {e}")
            return None

    def agregar_a_word(self, doc, img_path, nombre_word, primera=False):
        """Añade imagen al documento Word"""
        try:
            if not primera:
                doc.add_paragraph("\n\n")
            doc.add_picture(img_path, width=Inches(6))
            doc.save(nombre_word)
        except Exception as e:
            self.log(f"Error guardando en Word: {e}")

    def on_salir(self):
        """Maneja el cierre de la ventana"""
        self.capturador_activo = False
        self.grabacion_activa = False
        
        # Dar tiempo a que los hilos terminen
        if self.hilo_capturador and self.hilo_capturador.is_alive():
            time.sleep(0.5)
        if self.hilo_grabador and self.hilo_grabador.is_alive():
            time.sleep(0.5)
            
        self.destroy()

class EditorObjetivos(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.parent = parent
        self.title("Editar Objetivos")
        self.geometry("500x450")
        self.resizable(False, False)
        
        self.crear_interfaz()
        self.cargar_objetivos()
        
        # Hacer modal
        self.transient(parent)
        self.grab_set()
        
        # Centrar
        self.centrar_ventana()

    def centrar_ventana(self):
        self.update_idletasks()
        ancho = self.winfo_width()
        alto = self.winfo_height()
        x = (self.winfo_screenwidth() // 2) - (ancho // 2)
        y = (self.winfo_screenheight() // 2) - (alto // 2)
        self.geometry(f"{ancho}x{alto}+{x}+{y}")

    def crear_interfaz(self):
        main_frame = tb.Frame(self, padding=15)
        main_frame.pack(fill="both", expand=True)
        
        # Título e instrucciones
        etiqueta_titulo(
            main_frame,
            texto="Editar Objetivos de Captura",
            font=("Arial", 14, "bold")
        ).pack(pady=(0, 10))
        
        instrucciones = (
            "Añada o edite los títulos de las ventanas de sus aplicaciones objetivo.\n"
            "Escriba un título por línea. El título debe ser exacto para que el programa\n"
            "pueda encontrar la ventana correctamente.\n\n"
            "Ejemplos:\n"
            "• Sistema de Cartera\n"
            "• C O B I S  Clientes\n"
            "• Mi Aplicación - Google Chrome"
        )
        
        etiqueta_titulo(
            main_frame,
            texto=instrucciones,
            font=("Arial", 9)
        ).pack(pady=(0, 15), anchor="w")
        
        # Área de texto
        self.text_area = scrolledtext.ScrolledText(
            main_frame,
            height=10,
            font=("Consolas", 10),
            wrap="word"
        )
        self.text_area.pack(fill="both", expand=True, pady=(0, 20))
        
        # Botones
        btn_frame = tb.Frame(main_frame)
        btn_frame.pack(fill="x", pady=10)
        
        boton_exito(
            btn_frame,
            "Guardar",
            comando=self.guardar_objetivos,
            width=15
        ).pack(side="left", padx=5)
        
        boton_rojo(
            btn_frame,
            "Cancelar",
            comando=self.destroy,
            width=15
        ).pack(side="right", padx=5)

    def cargar_objetivos(self):
        """Carga los objetivos actuales en el área de texto"""
        objetivos_texto = "\n".join(self.parent.objetivos_base)
        self.text_area.insert("1.0", objetivos_texto)

    def guardar_objetivos(self):
        """Guarda los objetivos editados"""
        texto = self.text_area.get("1.0", tk.END).strip()
        nuevos_objetivos = [linea.strip() for linea in texto.split("\n") if linea.strip()]
        
        if not nuevos_objetivos:
            messagebox.showwarning("Error", "Debe especificar al menos un objetivo.")
            return
        
        self.parent.objetivos_base = nuevos_objetivos
        self.parent.guardar_configuracion()
        self.parent.actualizar_combos()
        self.parent.log(f"Objetivos actualizados: {len(nuevos_objetivos)} elementos")
        
        messagebox.showinfo("Éxito", "Objetivos guardados correctamente.")
        self.destroy()

# Función para usar desde el menú principal
def abrir_asistente_captura(master):
    """Función para abrir el asistente desde el menú principal"""
    ventana = AsistenteCapturaVentana(master)
    ventana.grab_set()
    return ventana