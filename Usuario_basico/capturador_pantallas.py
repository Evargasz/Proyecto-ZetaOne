import pyautogui
import keyboard
import time
import os
import json
from docx import Document
from docx.shared import Inches
from pywinauto import Desktop, uia_defines
from PIL import Image, ImageDraw
import math
from io import BytesIO
import win32clipboard
import win32con # Added for window style constants

# Configuración de aplicaciones desde config.json
def cargar_objetivos_configurados():
    """Carga los objetivos desde config.json"""
    try:
        config_file = "json\config.json"
        if os.path.exists(config_file):
            with open(config_file, 'r', encoding='utf-8') as f:
                config = json.load(f)
                return config.get("objetivos", [])
    except Exception as e:
        print(f"Error cargando configuración: {e}")
    return ["Sistema de Cartera", "C O B I S  Clientes", "C.O.B.I.S TERMINAL ADMINISTRATIVA"]

objetivos = cargar_objetivos_configurados()
APLICACIONES_CONOCIDAS = {}
for i, objetivo in enumerate(objetivos, 1):
    APLICACIONES_CONOCIDAS[str(i)] = objetivo
if len(objetivos) >= 2:
    APLICACIONES_CONOCIDAS[str(len(objetivos) + 1)] = f"{objetivos[0]} y {objetivos[1]}"
APLICACIONES_CONOCIDAS[str(len(APLICACIONES_CONOCIDAS) + 1)] = "Todas (Pantalla Completa)"

def seleccionar_aplicacion():
    """Muestra un menú para que el usuario elija qué aplicación capturar."""
    print("Selecciona la aplicación de la que deseas tomar capturas:")
    for key, value in APLICACIONES_CONOCIDAS.items():
        print(f"  {key}) {value}")
    print("  *) O escribe el título exacto de otra ventana")
    while True:
        choice = input("> ").strip()
        if choice in APLICACIONES_CONOCIDAS:
            valor = APLICACIONES_CONOCIDAS[choice]
            if " y " in valor and valor != "Todas (Pantalla Completa)":
                objetivos = cargar_objetivos_configurados()
                if len(objetivos) >= 2:
                    return [objetivos[0], objetivos[1]]
            return valor
        if choice:
            return choice
        print("Selección inválida. Por favor, elige un número o escribe un título.")

def encontrar_ventana(window_title):
    """Encuentra ventana usando win32gui (sin pywinauto)"""
    import win32gui
    def callback(hwnd, ventanas):
        if win32gui.IsWindowVisible(hwnd):
            texto = win32gui.GetWindowText(hwnd)
            if window_title.lower() in texto.lower() and texto.strip():
                ventanas.append(hwnd)
        return True
    ventanas = []
    try:
        win32gui.EnumWindows(callback, ventanas)
        if ventanas:
            hwnd = ventanas[0]
            class VentanaSimulada:
                def __init__(self, hwnd):
                    self.hwnd = hwnd
                def exists(self):
                    return win32gui.IsWindow(self.hwnd)
                def is_minimized(self):
                    return win32gui.IsIconic(self.hwnd)
                def restore(self):
                    win32gui.ShowWindow(self.hwnd, 9)
                def set_focus(self):
                    win32gui.SetForegroundWindow(self.hwnd)
                def rectangle(self):
                    rect = win32gui.GetWindowRect(self.hwnd)
                    class Rect:
                        def __init__(self, rect):
                            self.left = rect[0]
                            self.top = rect[1]
                            self.right = rect[2]
                            self.bottom = rect[3]
                        def width(self):
                            return self.right - self.left
                        def height(self):
                            return self.bottom - self.top
                    return Rect(rect)
            return VentanaSimulada(hwnd)
    except:
        pass
    return None

def encontrar_ventana_incrustada_bajo_cursor():
    """
    Encuentra la ventana "incrustada" bajo el cursor.
    Sube por la jerarquía desde el control bajo el cursor y se detiene en la
    primera ventana que tenga una barra de título (caption), o en la de más alto nivel.
    """
    try:
        import win32gui
        import win32api

        cursor_pos = win32api.GetCursorPos()
        hwnd = win32gui.WindowFromPoint(cursor_pos)
        
        if not hwnd:
            return None

        # Bucle para subir en la jerarquía
        while True:
            style = win32gui.GetWindowLong(hwnd, win32con.GWL_STYLE)
            if style & win32con.WS_CAPTION:
                break 

            parent = win32gui.GetParent(hwnd)
            if parent == 0:
                break
            
            hwnd = parent

        class VentanaSimulada:
            def __init__(self, hwnd):
                self.hwnd = hwnd
            def exists(self):
                return win32gui.IsWindow(self.hwnd)
            def is_minimized(self):
                return win32gui.IsIconic(self.hwnd)
            def restore(self):
                win32gui.ShowWindow(self.hwnd, 9)
            def set_focus(self):
                try:
                    win32gui.SetForegroundWindow(self.hwnd)
                except:
                    pass
            def rectangle(self):
                rect = win32gui.GetWindowRect(self.hwnd)
                class Rect:
                    def __init__(self, rect):
                        self.left = rect[0]
                        self.top = rect[1]
                        self.right = rect[2]
                        self.bottom = rect[3]
                    def width(self):
                        return self.right - self.left
                    def height(self):
                        return self.bottom - self.top
                return Rect(rect)
        
        return VentanaSimulada(hwnd)
        
    except Exception as e:
        print(f"Error encontrando ventana incrustada: {e}")
        return None

def agregar_a_word(doc, img_path, nombre_word, primera=False):
    if not primera:
        doc.add_paragraph("\n\n")
    doc.add_picture(img_path, width=Inches(6))
    doc.save(nombre_word)

def detectar_elemento_bajo_cursor():
    """Detecta el elemento UI más preciso bajo el cursor del mouse"""
    try:
        import win32gui
        import win32api
        cursor_x, cursor_y = win32api.GetCursorPos()
        hwnd = win32gui.WindowFromPoint((cursor_x, cursor_y))
        if hwnd:
            rect = win32gui.GetWindowRect(hwnd)
            width = rect[2] - rect[0]
            height = rect[3] - rect[1]
            if 10 < width < 800 and 10 < height < 200:
                return {
                    'left': rect[0],
                    'top': rect[1], 
                    'right': rect[2],
                    'bottom': rect[3],
                    'method': 'WindowFromPoint'
                }
        return None
    except Exception as e:
        print(f"Error detectando elemento: {e}")
        return None

def capturar_y_guardar(seleccion, carpeta_capturas, contador):
    """    Lógica centralizada para tomar una captura de pantalla con marcación precisa.
    Devuelve (nombre_archivo, screenshot_object) o None si falla.
    """
    try:
        screenshot = None
        rect_offset_x, rect_offset_y = 0, 0
        ventana_a_capturar = None

        if seleccion == "AUTODETECT":
            ventana_a_capturar = encontrar_ventana_incrustada_bajo_cursor()
            if not ventana_a_capturar:
                print("❌ No se pudo detectar una ventana bajo el cursor.")
                return None
        elif isinstance(seleccion, list):
            ventanas_encontradas = [encontrar_ventana(titulo) for titulo in seleccion]
            ventanas_validas = [v for v in ventanas_encontradas if v and v.exists()]
            if not ventanas_validas:
                ventanas_str = " y ".join(seleccion)
                print(f"❌ No encuentra la ventana o pantalla '{ventanas_str}' parametrizada abierta!")
                return None
            min_left = min(v.rectangle().left for v in ventanas_validas)
            min_top = min(v.rectangle().top for v in ventanas_validas)
            max_right = max(v.rectangle().right for v in ventanas_validas)
            max_bottom = max(v.rectangle().bottom for v in ventanas_validas)
            region_captura = (min_left, min_top, max_right - min_left, max_bottom - min_top)
            screenshot = pyautogui.screenshot(region=region_captura)
            rect_offset_x, rect_offset_y = min_left, min_top
        elif seleccion == "Todas (Pantalla Completa)":
            screenshot = pyautogui.screenshot()
        else:
            ventana_a_capturar = encontrar_ventana(seleccion)
            if not (ventana_a_capturar and ventana_a_capturar.exists()):
                print(f"❌ No encuentra la ventana o pantalla '{seleccion}' parametrizada abierta!")
                return None

        if screenshot is None and ventana_a_capturar:
            if ventana_a_capturar.is_minimized():
                ventana_a_capturar.restore()
            ventana_a_capturar.set_focus()
            time.sleep(0.2)
            window_rect = ventana_a_capturar.rectangle()
            if window_rect.width() <= 0 or window_rect.height() <= 0:
                print(f"La ventana no tiene un tamaño visible.")
                return None
            screenshot = pyautogui.screenshot(region=(window_rect.left, window_rect.top, window_rect.width(), window_rect.height()))
            rect_offset_x, rect_offset_y = window_rect.left, window_rect.top

        if screenshot is None:
            print("❌ Error desconocido: no se pudo tomar la captura.")
            return None

        elemento_detectado = detectar_elemento_bajo_cursor()
        if elemento_detectado:
            relative_left = max(0, elemento_detectado['left'] - rect_offset_x)
            relative_top = max(0, elemento_detectado['top'] - rect_offset_y)
            relative_right = min(screenshot.width, elemento_detectado['right'] - rect_offset_x)
            relative_bottom = min(screenshot.height, elemento_detectado['bottom'] - rect_offset_x)
            
            draw = ImageDraw.Draw(screenshot)
            draw.rectangle(
                [relative_left, relative_top, relative_right, relative_bottom], 
                outline="yellow", 
                width=4
            )
            draw.rectangle(
                [relative_left+1, relative_top+1, relative_right-1, relative_bottom-1], 
                outline="black", 
                width=1
            )
            print(f"✓ Elemento detectado y marcado.")
        else:
            cursor_x, cursor_y = pyautogui.position()
            relative_x = cursor_x - rect_offset_x
            relative_y = cursor_y - rect_offset_y
            
            if 0 <= relative_x < screenshot.width and 0 <= relative_y < screenshot.height:
                draw = ImageDraw.Draw(screenshot)
                flecha(draw, relative_x, relative_y, color="yellow", largo=40, grosor=3)
                print("✓ No se detectó un elemento, se dibujó una flecha en la posición del cursor.")

        nombre_archivo = os.path.join(carpeta_capturas, f"captura_{contador:03d}.png")
        screenshot.save(nombre_archivo)
        os.startfile(nombre_archivo)
        copy_clipboard(screenshot)
        return nombre_archivo, screenshot
    except Exception as e:
        print(f"Error durante la captura: {e}")
        return None

def flecha(draw, x, y, largo=50, angulo_grados=135, color="yellow", grosor=3):
    """
    Dibuja una flecha apuntando a las coordenadas (x, y).
    La flecha se origina desde una distancia 'largo' en la dirección 'angulo_grados'.
    """
    angulo_rad = math.radians(angulo_grados)
    start_x = x + largo * math.cos(angulo_rad)
    start_y = y + largo * math.sin(angulo_rad)
    draw.line([(start_x, start_y), (x, y)], fill=color, width=grosor)
    longitud_cabeza = 15
    angulo_cabeza = math.radians(30)
    linea_rad = math.atan2(y - start_y, x - start_x)
    p1_x = x - longitud_cabeza * math.cos(linea_rad - angulo_cabeza)
    p1_y = y - longitud_cabeza * math.sin(linea_rad - angulo_cabeza)
    p2_x = x - longitud_cabeza * math.cos(linea_rad + angulo_cabeza)
    p2_y = y - longitud_cabeza * math.sin(linea_rad + angulo_cabeza)
    draw.line([(x, y), (p1_x, p1_y)], fill=color, width=grosor)
    draw.line([(x, y), (p2_x, p2_y)], fill=color, width=grosor)

def copy_clipboard(img):
    output = BytesIO()
    img.convert("RGB").save(output, "BMP")
    data = output.getvalue()[14:]
    output.close()
    win32clipboard.OpenClipboard()
    win32clipboard.EmptyClipboard()
    win32clipboard.SetClipboardData(win32clipboard.CF_DIB, data)
    win32clipboard.CloseClipboard()

def main():
    if not APLICACIONES_CONOCIDAS:
        print("ERROR: No hay objetivos configurados en config.json")
        return
    seleccion = seleccionar_aplicacion()
    if not seleccion:
        return
    window_title = " y ".join(seleccion) if isinstance(seleccion, list) else seleccion
    print(f"\nListo para capturar '{window_title}'.")
    nombre_base = window_title.lower().replace(" ", "_").replace(".", "")
    carpeta_capturas = f"C:\\ZetaOne\\Capturas\\capturas_{nombre_base}"
    nombre_word = f"C:\\ZetaOne\\Capturas\\pantallazos_{nombre_base}.docx"
    os.makedirs(carpeta_capturas, exist_ok=True)
    print("Presiona F8 para capturar la ventana. Presiona ESC para salir.")
    if os.path.exists(nombre_word):
        doc = Document(nombre_word)
        primera_captura = False
    else:
        doc = Document()
        primera_captura = True
    contador = 1
    while True:
        if keyboard.is_pressed('f8'):
            resultado = capturar_y_guardar(seleccion, carpeta_capturas, contador)
            if resultado:
                nombre_archivo, _ = resultado
                print(f"¡Captura {contador} guardada en {nombre_archivo}!")
                agregar_a_word(doc, nombre_archivo, nombre_word, primera=primera_captura)
                primera_captura = False
                contador += 1
            time.sleep(0.8)
        if keyboard.is_pressed('esc'):
            print("¡Capturador finalizado!")
            break

def main_con_objetivo(objetivo_preseleccionado):
    """Función principal con objetivo preseleccionado desde la interfaz"""
    seleccion = objetivo_preseleccionado
    if seleccion == "AUTODETECT":
        window_title = "Ventana bajo el cursor"
    elif " y " in seleccion and seleccion != "Todas (Pantalla Completa)":
        objetivos = cargar_objetivos_configurados()
        if len(objetivos) >= 2:
            seleccion = [objetivos[0], objetivos[1]]
            window_title = " y ".join(seleccion)
    else:
        window_title = seleccion
    print(f"\n=== CAPTURADOR INICIADO ===")
    print(f"Objetivo: {window_title}")
    print("F8 = Capturar | ESC = Salir")
    print("===============================")
    nombre_base = "capturas_autodetect" if seleccion == "AUTODETECT" else window_title.lower().replace(" ", "_").replace(".", "")
    carpeta_capturas = f"C:\\ZetaOne\\Capturas\\{nombre_base}"
    nombre_word = f"C:\\ZetaOne\\Capturas\\pantallazos_{nombre_base}.docx"
    os.makedirs(carpeta_capturas, exist_ok=True)
    if os.path.exists(nombre_word):
        doc = Document(nombre_word)
        primera_captura = False
    else:
        doc = Document()
        primera_captura = True
    contador = 1
    while True:
        if keyboard.is_pressed('f8'):
            resultado = capturar_y_guardar(seleccion, carpeta_capturas, contador)
            if resultado:
                nombre_archivo, _ = resultado
                print(f"✓ Captura {contador} guardada: {os.path.basename(nombre_archivo)}")
                agregar_a_word(doc, nombre_archivo, nombre_word, primera=primera_captura)
                primera_captura = False
                contador += 1
            time.sleep(0.8)
        if keyboard.is_pressed('esc'):
            print("\n¡Capturador finalizado!")
            break

if __name__ == "__main__":
    main()
